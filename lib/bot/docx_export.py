# lib/bot/docx_export.py
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, Optional, Sequence
import datetime as dt

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None  # type: ignore
    Pt = None        # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore


def docx_available() -> bool:
    return Document is not None


def _now_jst_label() -> str:
    JST = dt.timezone(dt.timedelta(hours=9), name="Asia/Tokyo")
    return dt.datetime.now(JST).strftime("%Y-%m-%d %H:%M:%S %Z")


def build_meta_doc(
    *,
    user: str,
    chat_model: str,
    detail_label: str,
    detail: str,
    max_tokens: int,
    top_k: int,
    ts_jst: Optional[str] = None,
) -> Dict[str, Any]:
    return {
        "user": user,
        "chat_model": chat_model,
        "detail_label": detail_label,
        "detail": detail,
        "max_tokens": int(max_tokens),
        "top_k": int(top_k),
        "ts_jst": ts_jst or _now_jst_label(),
    }


def build_filters_doc(
    *,
    years_sel: Sequence[int] | None = None,
    pnos_sel_norm: Sequence[str] | None = None,
    shards: Sequence[str] | None = None,
) -> Dict[str, Any]:
    years = sorted(list(years_sel or []))
    pnos = sorted(list(pnos_sel_norm or []))
    shs = sorted(list(shards or []))
    return {"years": years, "pnos": pnos, "shards": shs}


def build_bot_answer_docx(
    prompt_text: str,
    answer_text: str,
    meta: Dict[str, Any],
    filters: Optional[Dict[str, Any]] = None,
) -> bytes:
    """
    質問＋回答＋メタ情報＋フィルタ情報を Word（.docx）に書き出す。
    戻り値は docx バイナリ。
    """
    if Document is None:
        raise RuntimeError("python-docx が見つかりません。`pip install python-docx` を実行してください。")

    doc = Document()

    title = doc.add_paragraph("Internal Bot 応答")
    if title.runs:
        title.runs[0].font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    m = doc.add_paragraph()
    m.add_run("Meta").bold = True
    doc.add_paragraph(f"User: {meta.get('user') or '(anonymous)'}")
    doc.add_paragraph(f"Model: {meta.get('chat_model')}")
    doc.add_paragraph(f"Detail: {meta.get('detail_label')} ({meta.get('detail')})")
    doc.add_paragraph(f"Max Tokens: {meta.get('max_tokens')}")
    doc.add_paragraph(f"Top-K: {meta.get('top_k')}")
    doc.add_paragraph(f"Generated At: {meta.get('ts_jst')}")

    if filters and any([filters.get("years"), filters.get("pnos"), filters.get("shards")]):
        doc.add_paragraph("")
        f_hdr = doc.add_paragraph("Filters")
        if f_hdr.runs:
            f_hdr.runs[0].bold = True

        if filters.get("years"):
            doc.add_paragraph(f"year: {', '.join(map(str, filters['years']))}")
        if filters.get("pnos"):
            doc.add_paragraph(f"pno: {', '.join(filters['pnos'])}")
        if filters.get("shards"):
            doc.add_paragraph(f"shards: {', '.join(filters['shards'])}")

    doc.add_paragraph("")
    p_hdr = doc.add_paragraph("質問（ユーザープロンプト）")
    if p_hdr.runs:
        p_hdr.runs[0].bold = True
    for ln in (prompt_text or "").splitlines():
        doc.add_paragraph(ln)

    doc.add_paragraph("")
    a_hdr = doc.add_paragraph("回答")
    if a_hdr.runs:
        a_hdr.runs[0].bold = True
    for ln in (answer_text or "").splitlines():
        doc.add_paragraph(ln)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def make_bot_answer_docx_bytes(
    *,
    prompt_text: str,
    answer_text: str,
    user: str,
    chat_model: str,
    detail_label: str,
    detail: str,
    max_tokens: int,
    top_k: int,
    years_sel: Sequence[int] | None = None,
    pnos_sel_norm: Sequence[str] | None = None,
    shards: Sequence[str] | None = None,
    ts_jst: Optional[str] = None,
) -> bytes:
    """
    pages 側で meta_doc / filters_doc を作らずに済む統合関数。
    """
    meta = build_meta_doc(
        user=user,
        chat_model=chat_model,
        detail_label=detail_label,
        detail=detail,
        max_tokens=max_tokens,
        top_k=top_k,
        ts_jst=ts_jst,
    )
    filters = build_filters_doc(
        years_sel=years_sel,
        pnos_sel_norm=pnos_sel_norm,
        shards=shards,
    )
    return build_bot_answer_docx(prompt_text, answer_text, meta, filters)

def make_default_docx_filename(*, prefix: str = "bot_answer", ts: Optional[dt.datetime] = None) -> str:
    """
    例: bot_answer_20251213_091530.docx
    """
    JST = dt.timezone(dt.timedelta(hours=9), name="Asia/Tokyo")
    now = ts or dt.datetime.now(JST)
    return f"{prefix}_{now:%Y%m%d_%H%M%S}.docx"


