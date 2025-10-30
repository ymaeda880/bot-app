# pages/10_ボット.py（実行タイムライン削除・合計コストのみ表示・シャードUI削除版）
from __future__ import annotations
from pathlib import Path
from typing import List, Tuple, Dict, Any, Set
import heapq
from itertools import count
import datetime as dt
import time
import json

import streamlit as st
import numpy as np
import sys

from config.path_config import PATHS
from config.sample_questions import SAMPLES2   # ★ サンプルは SAMPLES2 から
from lib.text_normalize import normalize_ja_text
from lib.prompts.bot_prompt import build_prompt
from lib.gpt_responder import GPTResponder, CompletionResult
from lib.rag.rag_utils import EmbeddingStore, NumpyVectorDB
from lib.costs import (
    estimate_chat_cost, estimate_embedding_cost, usd_to_jpy, DEFAULT_USDJPY, ChatUsage
)
from lib.openai_utils import count_tokens
from lib.bot_utils import (
    list_shard_dirs_openai, fmt_source, enrich_citations, citation_tag,
    parse_years, parse_pnos, norm_pno_forms, year_ok, pno_ok,
    scan_candidate_files, filters_caption,
)

_THIS = Path(__file__).resolve()
PROJECTS_ROOT = _THIS.parents[3]  # 3階層上 = projects/
if str(PROJECTS_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECTS_ROOT))

from common_lib.auth.auth_helpers import get_current_user_from_session_or_cookie
from common_lib.logs.jsonl_logger import JsonlLogger, sha256_short

from io import BytesIO
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None  # python-docx 未導入時のフォールバック用

# --- sys.path 調整（common_lib へ到達） ---
import sys
_THIS = Path(__file__).resolve()
APP_DIR = _THIS.parents[1]        # .../bot_app
PROJ_DIR = _THIS.parents[2]       # .../bot_project
MONO_ROOT = _THIS.parents[3]      # .../projects
for p in (MONO_ROOT, PROJ_DIR, APP_DIR):
    if str(p) not in sys.path:
        sys.path.insert(0, str(p))

# --- ロガー ---
_APP_DIR = Path(__file__).resolve().parents[1]
_PAGE_NAME = Path(__file__).stem
logger = JsonlLogger(app_dir=_APP_DIR, page_name=_PAGE_NAME)
INCLUDE_FULL_PROMPT_IN_LOG = True

VS_ROOT: Path = PATHS.vs_root
CHAT_MODEL = "gpt-5-mini"  # ★ モデル固定

# ===== Word 出力 =====
def _build_docx(prompt_text: str, answer_text: str, meta: Dict[str, Any], filters: Dict[str, Any] | None = None) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx が見つかりません。`pip install python-docx` を実行してください。")
    doc = Document()

    title = doc.add_paragraph("Internal Bot 応答")
    title.runs[0].font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    m = doc.add_paragraph()
    m.add_run("Meta").bold = True
    doc.add_paragraph(f"User: {meta.get('user') or '(anonymous)'}")
    doc.add_paragraph(f"Model: {meta.get('chat_model')}")
    doc.add_paragraph(f"Detail: {meta.get('detail_label')} ({meta.get('detail')})")
    doc.add_paragraph(f"Max Tokens: {meta.get('max_tokens')}")
    doc.add_paragraph(f"Top-K: {meta.get('top_k')}")
    doc.add_paragraph(f"Generated At: {meta.get('ts_jst')}")

    if filters and any([filters.get("years"), filters.get("pnos"), filters.get("shards")]):
        doc.add_paragraph("")
        f_hdr = doc.add_paragraph("Filters"); f_hdr.runs[0].bold = True
        if filters.get("years"):
            doc.add_paragraph(f"year: {', '.join(map(str, filters['years']))}")
        if filters.get("pnos"):
            doc.add_paragraph(f"pno: {', '.join(filters['pnos'])}")
        if filters.get("shards"):
            doc.add_paragraph(f"shards: {', '.join(filters['shards'])}")

    doc.add_paragraph("")
    p_hdr = doc.add_paragraph("質問（ユーザープロンプト）"); p_hdr.runs[0].bold = True
    for ln in (prompt_text or "").splitlines():
        doc.add_paragraph(ln)

    doc.add_paragraph("")
    a_hdr = doc.add_paragraph("回答"); a_hdr.runs[0].bold = True
    for ln in (answer_text or "").splitlines():
        doc.add_paragraph(ln)

    bio = BytesIO(); doc.save(bio)
    return bio.getvalue()


# ===== UI =====
st.set_page_config(page_title="Chat Bot (Sharded) — 簡略版", page_icon="💬", layout="wide")

col_title, col_user = st.columns([5, 2], vertical_alignment="center")
with col_title:
    st.title("💬 Internal Bot（簡略版）")
with col_user:
    current_user, _payload = get_current_user_from_session_or_cookie(st)
    if current_user:
        st.success(f"ログイン中: **{current_user}**")
    else:
        st.warning("未ログイン（Cookie 未検出）")

if "q" not in st.session_state:
    st.session_state.q = ""

def _set_q(x: str) -> None:
    st.session_state.q = x or ""

with st.expander("ℹ️ このページの使い方", expanded=False):
    st.markdown("""
### 1) 何ができる？
- 社内PDFをベクトルDBから **RAG 検索＋回答生成** します。
- **year / pno フィルタ** で範囲を絞れます。
- **モデルは固定（gpt-5-mini）**、**回答生成は常に OpenAI**、**出典は常に [S1] 表記**です。
- 生成した **質問＋回答** を **Word（.docx）** で保存できます（フィルタ情報つき）。

### 2) サイドバー（主な設定）
- **検索件数（Top-K）**、**詳しさ**、**最大出力トークン**、**System Instruction**、**表示モード（逐次/一括）**、**year / pno フィルタ**。
- 参照ファイルの限定や、解決パス/タイムライン表示は本ページでは省略されています。

### 3) 使い方の流れ
1. サイドバーで必要に応じて設定（year/pno など）  
2. 入力欄に質問を入れて **送信**  
3. 回答の下に **出典（[S1]…）** と **参照コンテキスト** を表示  
4. 必要に応じて **Word で保存** ボタンから .docx を出力
    """)

st.divider()

# ----- サイドバー（検索対象シャード UI は削除） -----
with st.sidebar:
    # st.page_link("app.py", label="🤖 appへ戻る")



    st.header("設定")

    chat_model = CHAT_MODEL  # 固定（UI 非表示）

    top_k = st.slider("検索件数（Top-K）", 1, 12, 6, 1)
    _detail_label = st.selectbox("詳しさ", ["簡潔", "標準", "詳細", "超詳細"], index=2)
    _detail_map = {"簡潔": "concise", "標準": "standard", "詳細": "detailed", "超詳細": "very_detailed"}
    detail = _detail_map[_detail_label]; detail_label = _detail_label

    max_tokens = st.slider("最大出力トークン", 1000, 40000, 12000, 500)
    # sys_inst = st.text_area("System Instruction", "あなたは優秀な社内のアシスタントです.", height=80)
    # st.text_area は削除
    sys_inst = "あなたは優秀な社内のアシスタントです."

    display_mode = st.radio("表示モード", ["逐次表示（ストリーム）", "一括表示"], index=0)

    st.divider(); st.subheader("year / pno フィルタ")
    years_input = st.text_input("year（カンマ区切り・任意）", value="", help="例: 2019,2023")
    pnos_input  = st.text_input("pno（プロジェクト番号, カンマ区切り・任意）", value="", help="例: 010,045,120")

    years_sel: Set[int] = parse_years(years_input)
    pnos_raw: Set[str]  = parse_pnos(pnos_input)
    pnos_sel_norm: Set[str] = set()
    for p in pnos_raw:
        pnos_sel_norm |= norm_pno_forms(p)

    st.caption(
        f"year フィルタ: {sorted(list(years_sel)) or '（未指定）'} / "
        f"pno フィルタ: {sorted(list(pnos_sel_norm)) or '（未指定）'}"
    )
    st.caption("※ 未入力なら全件対象。どちらか/両方を入力した場合のみフィルタします。")

    st.divider(); st.subheader("🧪 サンプル質問（SAMPLES2）")
    cat_options = ["（未選択）"] + list(SAMPLES2.keys())
    cat = st.selectbox("カテゴリを選択", cat_options, index=0)
    sample_options = ["（未選択）"] if cat == "（未選択）" else ["（未選択）"] + SAMPLES2.get(cat, [])
    sample = st.selectbox("サンプル質問を選択", sample_options, index=0)
    st.button("⬇️ この質問を入力欄へセット", use_container_width=True,
              disabled=(sample in ("", "（未選択）")), on_click=lambda: _set_q(sample))

# ----- 本文 -----
q = st.text_area("質問を入力", value=st.session_state.q, height=100,
                 placeholder="この社内ボットに質問してください…")
if q != st.session_state.q:
    st.session_state.q = q

go = st.button("送信", type="primary")

if go and st.session_state.q.strip():
    # ▼ ログ
    try:
        _prompt_text = st.session_state.q.strip()
        logger.append({
            "user": current_user or "(anonymous)",
            "action": "ask",
            "chat_model": chat_model,
            "detail_label": detail_label,
            "detail": detail,
            "cite": True,                 # 常に True（[S1] 促し）
            "max_tokens": int(max_tokens),
            "top_k": int(top_k),
            "preset": False,
            "prompt_hash": sha256_short(_prompt_text),
            **({"prompt": _prompt_text} if INCLUDE_FULL_PROMPT_IN_LOG else {}),
        })
    except Exception as _log_e:
        st.warning(f"ログ保存に失敗しました: {_log_e}")

    try:
        # ★ シャードは UI なしで「全シャード」を対象
        vs_backend_dir = PATHS.vs_root / "openai"
        if not vs_backend_dir.exists():
            st.warning(f"ベクトルが見つかりません（{vs_backend_dir}）。先にベクトル化を実行してください。")
            st.stop()

        all_ids = [p.name for p in list_shard_dirs_openai(PATHS.vs_root)]
        shard_dirs = [vs_backend_dir / s for s in all_ids]
        shard_dirs = [p for p in shard_dirs if p.is_dir() and (p / "vectors.npy").exists()]
        if not shard_dirs:
            st.warning("検索可能なシャードがありません。")
            st.stop()

        # 候補抽出（year/pno で前処理）
        cand_by_shard, cand_total = scan_candidate_files(shard_dirs, years_sel, pnos_sel_norm)
        if (years_sel or pnos_sel_norm) and cand_total == 0:
            st.warning("指定の year/pno に一致する**候補ファイルが 0 件**のため、検索を実行しません。条件を見直してください。")
            st.stop()
        st.caption(filters_caption(years_sel, pnos_sel_norm, cand_total if (years_sel or pnos_sel_norm) else None))

        # 埋め込み
        question = normalize_ja_text(st.session_state.q)
        estore = EmbeddingStore(backend="openai")
        embedding_tokens = count_tokens(st.session_state.q, "text-embedding-3-large")
        qv = estore.embed([question]).astype("float32")

        # 検索（TopK マージ）
        K = int(top_k)
        heap_: List[Tuple[float, int, int, Dict[str, Any]]] = []
        tie = count()
        for shp in shard_dirs:
            try:
                vdb = NumpyVectorDB(shp)
                local_k = max(K * 10, 50) if (years_sel or pnos_sel_norm) else K
                hits = vdb.search(qv, top_k=local_k, return_="similarity")
                for h in hits:
                    if isinstance(h, tuple) and len(h) == 3:
                        row_idx, score, meta = h
                    else:
                        score, meta = h
                        row_idx = -1
                    md = dict(meta or {}); md["shard_id"] = shp.name
                    if not year_ok(md, years_sel): continue
                    if not pno_ok(md, pnos_sel_norm): continue

                    sc = float(score)
                    if len(heap_) < K:
                        heapq.heappush(heap_, (sc, next(tie), row_idx, md))
                    elif sc > heap_[0][0]:
                        heapq.heapreplace(heap_, (sc, next(tie), row_idx, md))
            except Exception as e:
                st.warning(f"シャード {shp.name} の検索でエラー: {e}")

        raw_hits = [(rid, sc, md) for (sc, _tb, rid, md) in sorted(heap_, key=lambda x: x[0], reverse=True)]
        if not raw_hits:
            if years_sel or pnos_sel_norm:
                st.warning("条件には合う候補ファイルがありましたが、上位スコアに入りませんでした。\n"
                           "👉 Top-K を増やす／クエリを具体化する／シャードを絞る、のいずれかを試してください。")
            else:
                st.warning("該当コンテキストが見つかりませんでした。")
            st.stop()

        # 回答生成（OpenAI固定・出典促しON）
        labeled = [
            f"[S{i}] {m.get('text', '')}\n[meta: {fmt_source(m)} / score={float(s):.3f}]"
            for i, (_rid, s, m) in enumerate(raw_hits, 1)
        ]
        prompt = build_prompt(
            question,
            labeled,
            sys_inst=sys_inst,
            style_hint=detail,
            cite=True,          # 常に True
            strict=False
        )
        responder = GPTResponder()
        chat_prompt_tokens = chat_completion_tokens = 0

        if display_mode == "逐次表示（ストリーム）":
            st.subheader("🧠 回答（逐次表示）")
            with st.chat_message("assistant"):
                st.write_stream(
                    responder.stream(
                        model=chat_model,
                        system_instruction=sys_inst,
                        user_content=prompt,
                        max_output_tokens=int(max_tokens),
                        on_error_text="Responses stream error."
                    )
                )
            answer = responder.final_text or ""
            chat_prompt_tokens = responder.usage.input_tokens
            chat_completion_tokens = responder.usage.output_tokens
        else:
            st.subheader("🧠 回答（一括表示）")
            with st.spinner("回答生成中…"):
                result: CompletionResult = responder.complete(
                    model=chat_model,
                    system_instruction=sys_inst,
                    user_content=prompt,
                    max_output_tokens=int(max_tokens)
                )
            answer = result.text or ""
            chat_prompt_tokens = result.usage.input_tokens
            chat_completion_tokens = result.usage.output_tokens
            st.write(enrich_citations(answer, raw_hits))

        # 出典展開
        answer = enrich_citations(answer, raw_hits)
        import re as _re
        citations = _re.findall(r"\[S[^\]]+\]", answer)
        if citations:
            seen = []
            for c in citations:
                if c not in seen: seen.append(c)
            with st.expander("📝 出典拡張済み最終テキスト", expanded=False):
                st.caption("以下は回答内の出典タグを整理した一覧です。")
                st.markdown("### 📚 出典（出典ごとに改行）")
                st.text("\n".join(seen))

        # 参照コンテキスト
        with st.expander("🔍 参照コンテキスト（上位ヒット）", expanded=False):
            for i, (_rid, score, meta) in enumerate(raw_hits, 1):
                txt = str(meta.get("text", "") or "")
                label = fmt_source(meta)
                year = meta.get("year", "")
                pno = meta.get("pno", "") or meta.get("project_no", "")
                snippet = (txt[:1000] + "…") if len(txt) > 1000 else txt
                st.markdown(
                    f"**[{citation_tag(i, meta)}] score={float(score):.3f}**  "
                    f"`{label}` — **year:** {year} / **pno:** {pno}\n\n{snippet}"
                )

        # ===== Word ダウンロード =====
        try:
            JST = dt.timezone(dt.timedelta(hours=9), name="Asia/Tokyo")
            ts_jst = dt.datetime.now(JST).strftime("%Y-%m-%d %H:%M:%S %Z")
            meta_doc = {
                "user": current_user or "(anonymous)",
                "chat_model": chat_model,
                "detail_label": detail_label,
                "detail": detail,
                "max_tokens": int(max_tokens),
                "top_k": int(top_k),
                "ts_jst": ts_jst,
            }
            # シャードUIが無いので filters_doc の shards は空配列でOK
            filters_doc = {
                "years": sorted(list(years_sel)) if years_sel else [],
                "pnos": sorted(list(pnos_sel_norm)) if pnos_sel_norm else [],
                "shards": [],   # 本ページは全シャード固定
            }
            if Document is None:
                st.info("📄 Word 保存を有効にするには `pip install python-docx` を実行してください。")
            else:
                docx_bytes = _build_docx(st.session_state.q.strip(), answer or "", meta_doc, filters_doc)
                default_name = f"bot_answer_{dt.datetime.now(JST):%Y%m%d_%H%M%S}.docx"
                st.download_button(
                    "⬇️ プロンプト＋回答を Word で保存 (.docx)",
                    data=docx_bytes,
                    file_name=default_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        except Exception as _docx_e:
            st.warning(f"Word 保存でエラー: {_docx_e}")

      # ===== 合計コスト（簡易表示のみ）=====
        try:
            # 埋め込みコスト（dict: {"usd": float, "jpy": float}）
            emb_cost = estimate_embedding_cost("text-embedding-3-large", embedding_tokens)

            # チャットコスト（dict: {"usd": float, "jpy": float}）
            chat_cost = estimate_chat_cost(
                chat_model,
                ChatUsage(input_tokens=chat_prompt_tokens or 0, output_tokens=chat_completion_tokens or 0)
            )

            total_usd = float(emb_cost["usd"]) + float(chat_cost["usd"])
            total_jpy = float(emb_cost["jpy"]) + float(chat_cost["jpy"])

            st.info(f"📊 api使用料金の概算：**¥{total_jpy:,.2f}**（${total_usd:.4f}）")

        except Exception as _cost_e:
            st.caption(f"コスト計算に失敗しました: {_cost_e}")


    except Exception as e:
        st.error(f"検索/生成中にエラー: {e}")
else:
    st.info("質問を入力して『送信』を押してください。サイドバーで設定できます。")
