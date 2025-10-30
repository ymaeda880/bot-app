# pages/28_ボット（開発用）.py
# =============================================================================
# 💬 Internal Bot (RAG, Shards) — year/pno フィルタ付き（候補プリフィルタ＋TopKブースト）
# + ⏱ 実行タイミング計測（VectorDB 走査 / 埋め込み / 検索 / GPT API / ストリーム）
# + 👤 ログイン表示バッジ（右上）
# + 📝 JSONL ログ（ユーザー入力のみ / プリセット直送は除外 / 詳しさは日本語と英語コード）
# =============================================================================

from __future__ import annotations
from pathlib import Path
from typing import List, Tuple, Dict, Any, Iterable, Set
import heapq
from itertools import count
import json

import streamlit as st
import numpy as np

from config.path_config import PATHS
from config.sample_questions import SAMPLES, ALL_SAMPLES
from lib.text_normalize import normalize_ja_text
from lib.prompts.bot_prompt import build_prompt
from lib.gpt_responder import GPTResponder, CompletionResult
from lib.rag.rag_utils import EmbeddingStore, NumpyVectorDB
from lib.costs import (
    MODEL_PRICES_USD, EMBEDDING_PRICES_USD, DEFAULT_USDJPY,
    ChatUsage, estimate_chat_cost, estimate_embedding_cost, usd_to_jpy,
    render_usage_summary, _model_prices_per_1k
)
from lib.openai_utils import count_tokens
from lib.bot_utils import (
    list_shard_dirs_openai, norm_path, fmt_source,
    enrich_citations, citation_tag, get_openai_api_key,
    parse_inline_files, strip_inline_files,
    # 新規ユーティリティ
    to_halfwidth_digits, clean_pno_token, parse_years, parse_pnos, norm_pno_forms,
    year_ok, pno_ok, file_ok, scan_candidate_files, filters_caption,
)

from lib.metrics.timing_utils import Timings, stream_with_timing, render_metrics_ui
import datetime as dt
import time


from io import BytesIO
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None  # python-docx 未導入時のフォールバック用



# ============================================================
# sys.path に common_lib を確実に追加（上方探索）
# ============================================================
import sys
from pathlib import Path

_THIS = Path(__file__).resolve()
# 例: .../projects/bot_project/bot_app/pages/10_ボット.py
APP_DIR = _THIS.parents[1]        # .../bot_app
PROJ_DIR = _THIS.parents[2]       # .../bot_project
MONO_ROOT = _THIS.parents[3]      # .../projects  ← common_lib がここ直下にある想定

for p in (MONO_ROOT, PROJ_DIR, APP_DIR):
    if str(p) not in sys.path:
        sys.path.insert(0, str(p))
# ============================================================

# === 追加インポート（ログイン表示・JSONLロガー） ============================
from common_lib.auth.auth_helpers import get_current_user_from_session_or_cookie
from common_lib.logs.jsonl_logger import JsonlLogger, sha256_short

# アプリ／ページ情報とロガー初期化
_APP_DIR = Path(__file__).resolve().parents[1]
_PAGE_NAME = Path(__file__).stem
logger = JsonlLogger(app_dir=_APP_DIR, page_name=_PAGE_NAME)

# ユーザープロンプト本文をログ保存するか
INCLUDE_FULL_PROMPT_IN_LOG = True


# ===== 価格テーブル整形（USD/1K tok） =======================================
MODEL_PRICES_PER_1K = _model_prices_per_1k()

VS_ROOT: Path = PATHS.vs_root

# ===== wordへ出力するための関数 =======================================
def _build_docx(prompt_text: str, answer_text: str, meta: Dict[str, Any], filters: Dict[str, Any] | None = None) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx が見つかりません。`pip install python-docx` を実行してください。")
    doc = Document()

    # タイトル
    title = doc.add_paragraph("Internal Bot 応答")
    title.runs[0].font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # メタ情報
    doc.add_paragraph("")
    m = doc.add_paragraph()
    m.add_run("Meta").bold = True
    doc.add_paragraph(f"User: {meta.get('user') or '(anonymous)'}")
    doc.add_paragraph(f"Model: {meta.get('chat_model')}")
    doc.add_paragraph(f"Detail: {meta.get('detail_label')} ({meta.get('detail')})")
    doc.add_paragraph(f"Max Tokens: {meta.get('max_tokens')}")
    doc.add_paragraph(f"Top-K: {meta.get('top_k')}")
    doc.add_paragraph(f"Generated At: {meta.get('ts_jst')}")
  
    # フィルタ（指定があるときだけ）
    if filters:
        years = filters.get("years") or []
        pnos = filters.get("pnos") or []
        files = filters.get("file_whitelist") or []
        shards = filters.get("shards") or []

        # いずれかに要素があればセクション追加
        if any([years, pnos, files, shards]):
            doc.add_paragraph("")
            f_hdr = doc.add_paragraph("Filters")
            f_hdr.runs[0].bold = True

            if years:
                doc.add_paragraph(f"year: {', '.join(map(str, years))}")
            if pnos:
                doc.add_paragraph(f"pno: {', '.join(pnos)}")
            # if files:
            #     doc.add_paragraph("files:")
            #     for s in files[:200]:
            #         doc.add_paragraph(f" - {s}")
            #     if len(files) > 200:
            #         doc.add_paragraph(f" ... and {len(files)-200} more")
            if shards:
                doc.add_paragraph(f"shards: {', '.join(shards)}")






    # プロンプト
    doc.add_paragraph("")
    p_hdr = doc.add_paragraph("質問（ユーザープロンプト）")
    p_hdr.runs[0].bold = True
    for ln in (prompt_text or "").splitlines():
        doc.add_paragraph(ln)

    # 回答
    doc.add_paragraph("")
    a_hdr = doc.add_paragraph("回答")
    a_hdr.runs[0].bold = True
    for ln in (answer_text or "").splitlines():
        doc.add_paragraph(ln)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()



# ===== Streamlit 基本設定／タイトル ==========================================
st.set_page_config(page_title="Chat Bot (Sharded)", page_icon="💬", layout="wide")

# ▼▼ ログイン表示付きタイトル（右上にバッジ） ▼▼
col_title, col_user = st.columns([5, 2], vertical_alignment="center")
with col_title:
    st.title("💬 Internal Bot (RAG, Shards)")
with col_user:
    current_user, _payload = get_current_user_from_session_or_cookie(st)
    if current_user:
        st.success(f"ログイン中: **{current_user}**")
    else:
        st.warning("未ログイン（Cookie 未検出）")

if "q" not in st.session_state:
    st.session_state.q = ""

def _set_q(x: str) -> None:
    st.session_state.q = x or ""


# ▼ 使い方（ヘルプ）
with st.expander("ℹ️ このページの使い方（サイドバー設定を含む）", expanded=False):
    st.markdown("""
### 1) 何ができる？
- 社内PDFをシャーディングしたベクトルDBから **RAG 検索＋回答生成** を行います。
- **year / pno フィルタ**や**参照ファイル限定**で対象を絞り込み可能。
- 生成した **回答＋あなたの質問**を **Word（.docx）** で保存できます（フィルタ情報も出力）。

---

### 2) サイドバーの主な設定
- **モデル（Responses）**: 回答モデルを選択（例: `gpt-5-mini`）。
- **検索件数（Top-K）**: コンテキスト取得の件数。（適宜調整してください．）
- **詳しさ**: 出力の粒度（「簡潔/標準/詳細/超詳細」）。  
  ※ ログには **日本語ラベル**と**英語コード**の両方（例: `詳細` / `detailed`）が保存されます。
- **出典を [S1] で促す**: 回答に出典タグを含めるか。
- **最大出力トークン**: 回答の長さ上限。（これは変更の必要はありません．）
- **回答生成**:  （これは変更の必要はありません．）
  - `OpenAI`…検索＋要約（通常）  
  - `Retrieve-only`…**要約なし**でコンテキストのみ表示
- **System Instruction**: モデルへの前提指示（任意）。（これは変更の必要はありません．）
- **表示モード**: 逐次（ストリーム）or 一括。（これは変更の必要はありません．）
- **検索対象シャード**: どのシャードを検索するか（未選択=すべて）。
- **year / pno フィルタ**: 年やプロジェクト番号で対象を絞り込み。未入力なら全件。
- **参照ファイル（任意）**: `2025/foo.pdf, 2024/bar.pdf` のように **特定ファイルだけ**に限定したいときに指定。（これは変更の必要はありません．）
- **🧪 デモ用サンプル質問**:  
  - 「⬇️ この質問を入力欄へセット」で入力欄にコピー  
  - 「🚀 サンプルで即送信」で即実行  

> 参考: 下段の **解決パス** はパスの確認用です（編集はできません）。

---

### 3) 使い方の流れ
1. サイドバーを設定する（必要に応じて year/pno/ファイル限定も指定）。
2. 画面中央の「質問を入力」に聞きたい内容を入力。
3. **送信** を押す。  
   - 逐次表示を選んでいれば、回答がストリームで出ます。  
   - 回答下部の「🔍 参照コンテキスト」で、使われた上位文脈を確認できます。
4. 回答が出たら、**「⬇️ プロンプト＋回答を Word で保存」** ボタンで .docx を保存。  
   - Word には **質問・回答・メタ情報（モデル/詳しさ/Top-K等）**に加え、  
     指定した **year/pno/シャード/ファイル限定** のフィルタも記録されます。  

---

### 4) 画面各部の見方
- **🧠 回答**: 生成回答。出典タグ `[S1]` 等は後述の拡張で展開されます。
- **📝 出典拡張済み最終テキスト**: `[Sx]` をソース表記に展開した一覧。
- **🔍 参照コンテキスト**: 上位ヒットのスニペットとメタ（score, year, pno など）。
- **📊 使用量の概算**: 埋め込み/チャットのトークン・概算コスト。
- **🧠 メモリ状況（回答前/後）**: 実行前後のメモリの簡易メトリクス。

---

### 5) トラブルシューティング
- **「検索可能なシャードがありません」**: 先にベクトル化を実行してください（`vectors.npy` が必要）。
- **該当コンテキストが見つからない**: Top-K を増やす / **クエリを具体化（具体的な質問を行なってください）** / シャードを絞る。
    """)


####################
# メモリ監視（任意）
#####################
from lib.monitors.ui_memory_monitor import render_memory_kpi_row

st.divider()
st.markdown("### 🧠 メモリ状況（参考）")
render_memory_kpi_row()

# ===== サイドバー ==============================================================
with st.sidebar:
    st.header("設定")

    model_candidates = [m for m in MODEL_PRICES_PER_1K if m.startswith(("gpt-5", "gpt-4.1"))]
    all_models_sorted = sorted(model_candidates, key=lambda x: (0 if x.startswith("gpt-5") else 1, x))
    if not all_models_sorted:
        st.error("利用可能な Responses モデルが見つかりません。lib/costs.MODEL_PRICES_USD を確認してください。")
        st.stop()
    default_idx = all_models_sorted.index("gpt-5-mini") if "gpt-5-mini" in all_models_sorted else 0
    chat_model = st.selectbox("モデル（Responses）", all_models_sorted, index=default_idx)

    top_k = st.slider("検索件数（Top-K）", 1, 12, 6, 1)

    # ▼▼ 詳しさ：日本語ラベル保持＋英語コード変換 ▼▼
    _detail_label = st.selectbox("詳しさ", ["簡潔", "標準", "詳細", "超詳細"], index=2)
    _detail_map = {"簡潔": "concise", "標準": "standard", "詳細": "detailed", "超詳細": "very_detailed"}
    detail = _detail_map[_detail_label]     # 既存 build_prompt 用
    detail_label = _detail_label            # ログ用（日本語ラベル）

    cite = st.checkbox("出典を [S1] で促す", True)
    max_tokens = st.slider("最大出力トークン", 1000, 40000, 12000, 500)
    answer_backend = st.radio("回答生成", ["OpenAI", "Retrieve-only"], index=0)
    sys_inst = st.text_area("System Instruction", "あなたは優秀な社内のアシスタントです.", height=80)
    display_mode = st.radio("表示モード", ["逐次表示（ストリーム）", "一括表示"], index=0)

    st.divider(); st.subheader("検索対象シャード（OpenAI）")
    shard_dirs_all = list_shard_dirs_openai(VS_ROOT)
    shard_ids_all = [p.name for p in shard_dirs_all]
    target_shards = st.multiselect("（未選択=すべて）", shard_ids_all, default=shard_ids_all)

    # ✅ year/pno 入力（パースは bot_utils に委譲）
    st.divider(); st.subheader("year / pno フィルタ")
    years_input = st.text_input("year（カンマ区切り・任意）", value="", help="例: 2019,2023")
    pnos_input  = st.text_input("pno（プロジェクト番号, カンマ区切り・任意）", value="", help="例: 010,045,120")

    years_sel: Set[int] = parse_years(years_input)
    pnos_raw: Set[str]  = parse_pnos(pnos_input)

    # pno の表記ゆれ（元/ゼロ除去/3桁ゼロ埋め）を吸収
    pnos_sel_norm: Set[str] = set()
    for p in pnos_raw:
        pnos_sel_norm |= norm_pno_forms(p)

    st.caption(
        f"year フィルタ: {sorted(list(years_sel)) or '（未指定）'} / "
        f"pno フィルタ: {sorted(list(pnos_sel_norm)) or '（未指定）'}"
    )
    st.caption("※ 未入力なら全件対象。どちらか/両方を入力した場合のみフィルタします。")

    st.caption("参照ファイルを追加（例: 2025/foo.pdf, 2024/bar.pdf）")
    file_whitelist_str = st.text_input("参照ファイル（任意）", value="")
    ui_file_whitelist = {s.strip() for s in file_whitelist_str.split(",") if s.strip()}

    api_key = get_openai_api_key()
    if not api_key and answer_backend == "OpenAI":
        st.error("OpenAI APIキーが secrets.toml / 環境変数にありません。")

    st.divider(); st.markdown("### 📂 解決パス（参照用）")
    st.text_input("VS_ROOT", str(PATHS.vs_root), disabled=True)
    st.text_input("PDF_ROOT", str(PATHS.pdf_root), disabled=True)
    st.text_input("BACKUP_ROOT", str(PATHS.backup_root), disabled=True)
    if hasattr(PATHS, "data_root"):
        st.text_input("DATA_ROOT", str(PATHS.data_root), disabled=True)

    st.divider(); st.subheader("🧪 デモ用サンプル質問")
    cat = st.selectbox("カテゴリを選択", ["（未選択）"] + list(SAMPLES.keys()))
    sample_options = ["（未選択）"] if cat == "（未選択）" else ["（未選択）"] + SAMPLES.get(cat, [])
    sample = st.selectbox("サンプル質問を選択", sample_options, index=0)
    cols_demo = st.columns(2)
    with cols_demo[0]:
        st.button("⬇️ この質問を入力欄へセット", width='stretch',
                  disabled=(sample in ("", "（未選択）")), on_click=lambda: _set_q(sample))
    with cols_demo[1]:
        st.button("🎲 ランダム挿入", width='stretch',
                  on_click=lambda: _set_q(str(np.random.choice(ALL_SAMPLES)) if ALL_SAMPLES else ""))
    send_now = st.button("🚀 サンプルで即送信", width='stretch',
                         disabled=(st.session_state.q.strip() == ""))

    # ▼▼ プリセット直送（🚀）を判定（ログ抑制に使用） ▼▼
    is_preset_direct_send = bool(send_now)

    #### メモリ状況（回答前 / 回答後）
    st.divider()
    st.subheader("🧠 メモリ状況（回答前 / 回答後）")
    # 回答前/後のスナップショット描画先（この順序で下に並ぶ）
    mem_pre_box = st.container()
    mem_post_box = st.container()


# ===== 本文 =========================================================
q = st.text_area("質問を入力", value=st.session_state.q, height=100,
                 placeholder="この社内ボットに質問してください…")
if q != st.session_state.q:
    st.session_state.q = q

go = st.button("送信", type="primary")
go = go or bool(locals().get("send_now"))

# ===== 実行：検索 → 生成 → コスト → 参照 ===========================
if go and st.session_state.q.strip():
    timings = Timings()
    timings.mark("pipeline_start")

    # ▼▼ 回答前スナップショット ▼▼
    with mem_pre_box:
        st.caption("（回答前スナップショット）")
        render_memory_kpi_row()

    # ▼▼ ユーザー入力プロンプトのログ（🚀直送は除外） ▼▼
    try:
        if not is_preset_direct_send:
            _prompt_text = st.session_state.q.strip()
            logger.append({
                "user": current_user or "(anonymous)",
                "action": "ask",
                "chat_model": chat_model,
                "detail_label": detail_label,   # 例：「詳細」「超詳細」
                "detail": detail,               # 例："detailed"
                "cite": bool(cite),
                "max_tokens": int(max_tokens),
                "top_k": int(top_k),
                "preset": False,                # プリセット直送ではない
                "prompt_hash": sha256_short(_prompt_text),
                **({"prompt": _prompt_text} if INCLUDE_FULL_PROMPT_IN_LOG else {}),
            })
        # else: プリセット直送は保存しない
    except Exception as _log_e:
        st.warning(f"ログ保存に失敗しました: {_log_e}")

    try:
        vs_backend_dir = PATHS.vs_root / "openai"
        if not vs_backend_dir.exists():
            st.warning(f"ベクトルが見つかりません（{vs_backend_dir}）。先にベクトル化を実行してください。")
            st.stop()

        selected_ids = target_shards or [p.name for p in list_shard_dirs_openai(PATHS.vs_root)]
        shard_dirs = [vs_backend_dir / s for s in selected_ids]
        shard_dirs = [p for p in shard_dirs if p.is_dir() and (p / "vectors.npy").exists()]
        if not shard_dirs:
            st.warning("検索可能なシャードがありません。")
            st.stop()

        # インライン [[files: ...]] と UI ホワイトリストを合算（正規化は file_ok 内で再利用）
        inline_files = parse_inline_files(st.session_state.q)
        effective_whitelist_norm: Set[str] = {norm_path(x) for x in (ui_file_whitelist | inline_files)}

        # ---- meta.jsonl 走査（候補抽出） ----
        timings.mark("scan_start")
        cand_by_shard, cand_total = scan_candidate_files(shard_dirs, years_sel, pnos_sel_norm)
        timings.mark("scan_end")

        if (years_sel or pnos_sel_norm) and cand_total == 0:
            st.warning("指定の year/pno に一致する**候補ファイルが 0 件**のため、検索を実行しません。条件を見直してください。")
            st.stop()

        # year/pno 由来の候補をホワイトリストに合成
        for _sd, files in cand_by_shard.items():
            effective_whitelist_norm |= files

        # デバッグ表記
        st.caption(filters_caption(years_sel, pnos_sel_norm, cand_total if (years_sel or pnos_sel_norm) else None))

        # ---- クエリ埋め込み ----
        question = normalize_ja_text(strip_inline_files(st.session_state.q))
        estore = EmbeddingStore(backend="openai")
        question_tokens = count_tokens(st.session_state.q, "text-embedding-3-large")

        timings.mark("embed_start")
        qv = estore.embed([question]).astype("float32")
        timings.mark("embed_end")

        # ---- 各シャード TopK をマージ ----
        timings.mark("search_start")
        K = int(top_k)
        heap_: List[Tuple[float, int, int, Dict[str, Any]]] = []  # (score, tiebreak, row_idx, meta)
        tie = count()
        for shp in shard_dirs:
            try:
                vdb = NumpyVectorDB(shp)
                local_k = K
                if years_sel or pnos_sel_norm or effective_whitelist_norm:
                    local_k = max(K * 10, 50)
                hits = vdb.search(qv, top_k=local_k, return_="similarity")
                for h in hits:
                    if isinstance(h, tuple) and len(h) == 3:
                        row_idx, score, meta = h
                    else:
                        score, meta = h
                        row_idx = -1

                    md = dict(meta or {}); md["shard_id"] = shp.name

                    # 統一判定（year/pno/file）
                    if not year_ok(md, years_sel):                 # 年
                        continue
                    if not pno_ok(md, pnos_sel_norm):              # プロジェクト番号
                        continue
                    if not file_ok(md, effective_whitelist_norm):  # ファイル限定
                        continue

                    sc = float(score)
                    if len(heap_) < K:
                        heapq.heappush(heap_, (sc, next(tie), row_idx, md))
                    elif sc > heap_[0][0]:
                        heapq.heapreplace(heap_, (sc, next(tie), row_idx, md))
            except Exception as e:
                st.warning(f"シャード {shp.name} の検索でエラー: {e}")
        timings.mark("search_end")

        raw_hits = [(rid, sc, md) for (sc, _tb, rid, md) in sorted(heap_, key=lambda x: x[0], reverse=True)]
        if not raw_hits:
            if years_sel or pnos_sel_norm:
                st.warning("条件には合う候補ファイルがありましたが、上位スコアに入りませんでした。\n"
                           "👉 Top-K を増やす／クエリを具体化する／シャードを絞る、のいずれかを試してください。")
            else:
                st.warning("該当コンテキストが見つかりませんでした。")
            st.stop()

        # ---- 回答生成 ----
        chat_prompt_tokens = chat_completion_tokens = 0
        use_backend = "OpenAI" if (answer_backend == "OpenAI" and api_key) else "Retrieve-only"

        if use_backend == "OpenAI":
            labeled = [
                f"[S{i}] {m.get('text', '')}\n[meta: {fmt_source(m)} / score={float(s):.3f}]"
                for i, (_rid, s, m) in enumerate(raw_hits, 1)
            ]
            prompt = build_prompt(
                question, labeled, sys_inst=sys_inst, style_hint=detail, cite=cite, strict=False
            )
            responder = GPTResponder(api_key=api_key)
            use_stream = (display_mode == "逐次表示（ストリーム）")

            # 🔹 逐次出力（ストリーム）
            if use_stream:
                st.subheader("🧠 回答（逐次表示）")
                timings.mark("gpt_req_start")
                with st.chat_message("assistant"):
                    st.write_stream(
                        stream_with_timing(
                            responder.stream(
                                model=chat_model,
                                system_instruction=sys_inst,
                                user_content=prompt,
                                max_output_tokens=int(max_tokens),
                                on_error_text="Responses stream error."
                            ),
                            timings,
                            first_key="gpt_first_token",
                            done_key="gpt_done",
                        )
                    )
                answer = responder.final_text or ""
                chat_prompt_tokens = responder.usage.input_tokens
                chat_completion_tokens = responder.usage.output_tokens
                # 念のためのフォールバック
                timings.marks.setdefault("gpt_first_token", timings.marks.get("gpt_req_start"))
                timings.marks.setdefault("gpt_done", dt.datetime.now())
                timings.perf.setdefault("gpt_first_token", timings.perf.get("gpt_req_start", time.perf_counter()))
                timings.perf.setdefault("gpt_done", time.perf_counter())

            # 🔹 一括表示モード
            else:
                st.subheader("🧠 回答（一括表示）")
                timings.mark("gpt_req_start")
                with st.spinner("回答生成中…"):
                    result: CompletionResult = responder.complete(
                        model=chat_model,
                        system_instruction=sys_inst,
                        user_content=prompt,
                        max_output_tokens=int(max_tokens)
                    )
                # 一括は「最初のトークン＝完了」とみなす
                timings.mark("gpt_first_token")
                timings.mark("gpt_done")

                answer = result.text or ""
                chat_prompt_tokens = result.usage.input_tokens
                chat_completion_tokens = result.usage.output_tokens

                st.write(enrich_citations(answer, raw_hits))

            # ---- 出典処理（両モード共通）----
            answer = enrich_citations(answer, raw_hits)
            import re as _re
            citations = _re.findall(r"\[S[^\]]+\]", answer)
            if citations:
                seen = []
                for c in citations:
                    if c not in seen:
                        seen.append(c)
                with st.expander("📝 出典拡張済み最終テキスト", expanded=False):
                    st.caption("以下は回答内の出典タグを整理した一覧です。")
                    st.markdown("### 📚 出典（出典ごとに改行）")
                    st.text("\n".join(seen))
            else:
                st.caption("出典タグは検出されませんでした。")

        else:
            st.subheader("🧩 取得のみ（要約なし）")
            st.info("Retrieve-only モードです。下の参照コンテキストをご覧ください。")
            answer = ""  # 念のため

        
        # ===== ここから Word ダウンロード =====
        try:
            # 生成時刻（JST）
            JST = dt.timezone(dt.timedelta(hours=9), name="Asia/Tokyo")
            ts_jst = dt.datetime.now(JST).strftime("%Y-%m-%d %H:%M:%S %Z")

            # ドキュメント用メタ
            meta_doc = {
                "user": current_user or "(anonymous)",
                "chat_model": chat_model,
                "detail_label": detail_label,
                "detail": detail,
                "max_tokens": int(max_tokens),
                "top_k": int(top_k),
                "ts_jst": ts_jst,
            }

            # ✅ フィルタ情報を定義
            # filters_doc = {
            #     "years": sorted(list(years_sel)) if years_sel else [],
            #     "pnos": sorted(list(pnos_sel_norm)) if pnos_sel_norm else [],
            #     "file_whitelist": sorted(list(effective_whitelist_norm)) if effective_whitelist_norm else [],
            #     "shards": target_shards or [],
            # }

            # 明示指定のみ（UI + インライン）を記録したい場合
            explicit_files = sorted(list({norm_path(x) for x in (ui_file_whitelist | inline_files)}))
            filters_doc = {
                "years": sorted(list(years_sel)) if years_sel else [],
                "pnos": sorted(list(pnos_sel_norm)) if pnos_sel_norm else [],
                "file_whitelist": explicit_files,  # ← こちらに置き換え
                "shards": target_shards or [],
}

            prompt_text_for_doc = st.session_state.q.strip()
            answer_text_for_doc = answer or ""

            if Document is None:
                st.info("📄 Word 保存を有効にするには `pip install python-docx` を実行してください。")
            else:
                # ✅ filters_doc を渡す！
                docx_bytes = _build_docx(prompt_text_for_doc, answer_text_for_doc, meta_doc, filters_doc)
                default_name = f"bot_answer_{dt.datetime.now(JST):%Y%m%d_%H%M%S}.docx"
                st.download_button(
                    "⬇️ プロンプト＋回答を Word で保存 (.docx)",
                    data=docx_bytes,
                    file_name=default_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        except Exception as _docx_e:
            st.warning(f"Word 保存でエラー: {_docx_e}")
        # ===== ここまで Word ダウンロード =====




        # ---- 参照コンテキスト ----
        with st.expander("🔍 参照コンテキスト（上位ヒット）", expanded=False):
            for i, (_rid, score, meta) in enumerate(raw_hits, 1):
                txt = str(meta.get("text", "") or "")
                label = fmt_source(meta)
                year = meta.get("year", "")
                pno = meta.get("pno", "") or meta.get("project_no", "")
                snippet = (txt[:1000] + "…") if len(txt) > 1000 else txt
                st.markdown(
                    f"**[{citation_tag(i, meta)}] score={float(score):.3f}**  "
                    f"`{label}` — **year:** {year} / **pno:** {pno}\n\n{snippet}"
                )

        # ---- 使用量の概算 ----
        render_usage_summary(
            embedding_model="text-embedding-3-large",
            embedding_tokens=question_tokens,
            chat_model=chat_model,
            chat_prompt_tokens=chat_prompt_tokens,
            chat_completion_tokens=chat_completion_tokens,
            use_backend_openai=(use_backend == "OpenAI"),
            title="📊 使用量の概算",
        )
        # ⏱ パイプライン終了
        timings.mark("pipeline_end")
        render_metrics_ui(timings)

        # ⏱ 回答後スナップショット
        with mem_post_box:
            st.caption("（回答後スナップショット）")
            render_memory_kpi_row()

    except Exception as e:
        st.error(f"検索/生成中にエラー: {e}")
        # 例外時も可能な範囲でタイムラインを表示
        timings.marks.setdefault("pipeline_end", dt.datetime.now())
        timings.perf.setdefault("pipeline_end", time.perf_counter())
        render_metrics_ui(timings)
else:
    st.info("質問を入力して『送信』を押してください。サイドバーで設定できます。")





# 末尾：メモリ監視（任意で二重表示を避けるためコメントアウト例）
# st.divider()
# st.markdown("### 🧠 メモリ状況（参考）")
# render_memory_kpi_row()
